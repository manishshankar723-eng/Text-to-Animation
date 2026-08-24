"""
plan_export.py — take a content calendar out of the app.

A plan is only useful if the creator can work from it: in a spreadsheet they
sort and tick off, a document they share with a client, or a CSV they import
into Notion / Airtable / a scheduling tool. So the same plan exports three ways:

    .xlsx  — the working format. One row per upload, frozen header, filters on,
             column widths set so nothing needs resizing before it's usable.
    .docx  — the shareable format. Reads as a document, not a data dump.
    .csv   — the portable format. Stdlib only, imports anywhere.

Every exporter takes the same `plan` dict from plan_agent.generate_plan().
"""

import csv
import io
import logging

logger = logging.getLogger(__name__)

# Column order, shared by every format so the three exports agree. `slot` and
# `title` lead because that is what a creator scans for.
COLUMNS = [
    ("slot", "When"),
    ("title", "Title"),
    ("hook", "Hook (first 3 seconds)"),
    ("format", "Format"),
    ("pillar", "Pillar"),
    ("outline", "Outline"),
    ("keywords", "Keywords"),
    ("cta", "Call to action"),
    ("goal", "Goal"),
    ("effort", "Effort"),
]

# Roughly how wide each column wants to be, in characters.
_WIDTHS = {
    "slot": 16, "title": 52, "hook": 46, "format": 20, "pillar": 20,
    "outline": 60, "keywords": 32, "cta": 28, "goal": 12, "effort": 10,
}


def _cell(item: dict, key: str) -> str:
    """One field as flat text. Lists become readable lines, not Python reprs."""
    value = item.get(key, "")
    if isinstance(value, list):
        if key == "outline":
            return "\n".join(f"{i}. {v}" for i, v in enumerate(value, 1))
        return ", ".join(str(v) for v in value)
    return str(value or "")


def _title(plan: dict) -> str:
    return (plan.get("title") or "Content plan").strip() or "Content plan"


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------
def to_csv(plan: dict) -> bytes:
    """The calendar as CSV. UTF-8 BOM so Excel opens it without mangling."""
    buf = io.StringIO(newline="")
    writer = csv.writer(buf)
    writer.writerow([label for _, label in COLUMNS])
    for item in plan.get("items") or []:
        writer.writerow([_cell(item, key) for key, _ in COLUMNS])
    return buf.getvalue().encode("utf-8-sig")


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
def to_xlsx(plan: dict) -> bytes:
    """The calendar as a working spreadsheet.

    Two sheets: the calendar itself, and a Strategy sheet carrying the summary,
    pillars and assumptions — so the thinking travels with the schedule instead
    of being lost the moment it leaves the app.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Calendar"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="3B3F51")
    wrap = Alignment(vertical="top", wrap_text=True)

    ws.append([label for _, label in COLUMNS])
    for i, _ in enumerate(COLUMNS, start=1):
        c = ws.cell(row=1, column=i)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(vertical="center", wrap_text=True)

    for item in plan.get("items") or []:
        ws.append([_cell(item, key) for key, _ in COLUMNS])

    for i, (key, _) in enumerate(COLUMNS, start=1):
        ws.column_dimensions[get_column_letter(i)].width = _WIDTHS.get(key, 20)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = wrap

    # Freeze the header and switch filters on: the first two things anyone does
    # to a schedule like this by hand.
    ws.freeze_panes = "A2"
    if ws.max_row >= 1:
        ws.auto_filter.ref = f"A1:{get_column_letter(len(COLUMNS))}{ws.max_row}"

    strat = wb.create_sheet("Strategy")
    strat.column_dimensions["A"].width = 22
    strat.column_dimensions["B"].width = 100
    bold = Font(bold=True)

    def row(label: str, value: str) -> None:
        strat.append([label, value])
        strat.cell(row=strat.max_row, column=1).font = bold
        strat.cell(row=strat.max_row, column=2).alignment = wrap

    row("Plan", _title(plan))
    row("Covers", f"{plan.get('months', 1)} month(s)")
    row("Cadence", plan.get("cadence", ""))
    row("Uploads", str(len(plan.get("items") or [])))
    if plan.get("summary"):
        row("Strategy", plan["summary"])
    if plan.get("pillars"):
        strat.append([])
        strat.append(["Content pillars", ""])
        strat.cell(row=strat.max_row, column=1).font = bold
        for p in plan["pillars"]:
            strat.append([p.get("name", ""), p.get("why", "")])
            strat.cell(row=strat.max_row, column=2).alignment = wrap
    if plan.get("assumptions"):
        strat.append([])
        strat.append(["Assumptions", ""])
        strat.cell(row=strat.max_row, column=1).font = bold
        for a in plan["assumptions"]:
            strat.append(["", a])
            strat.cell(row=strat.max_row, column=2).alignment = wrap

    out = io.BytesIO()
    wb.save(out)
    logger.info("[plan-export] xlsx: %d item(s)", len(plan.get("items") or []))
    return out.getvalue()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def to_docx(plan: dict) -> bytes:
    """The calendar as a document to read or send on.

    Laid out per upload rather than as a giant table: a 40-row table with a
    paragraph in every cell is unreadable on paper, which is where a docx ends
    up. Each upload gets a heading and its details underneath.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(_title(plan), level=0)

    meta = f"{plan.get('months', 1)} month(s) · {plan.get('cadence', '')} · " \
           f"{len(plan.get('items') or [])} uploads"
    doc.add_paragraph(meta.strip(" ·"))

    if plan.get("summary"):
        doc.add_heading("Strategy", level=1)
        doc.add_paragraph(plan["summary"])

    if plan.get("pillars"):
        doc.add_heading("Content pillars", level=1)
        for p in plan["pillars"]:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(p.get("name", ""))
            run.bold = True
            if p.get("why"):
                para.add_run(f" — {p['why']}")

    doc.add_heading("Calendar", level=1)
    for i, item in enumerate(plan.get("items") or [], start=1):
        head = doc.add_heading(level=2)
        slot = item.get("slot", "")
        head.add_run(f"{i}. {item.get('title', '')}")
        if slot:
            sub = doc.add_paragraph()
            r = sub.add_run(slot)
            r.italic = True
            r.font.size = Pt(9)

        def field(label: str, value: str) -> None:
            if not value:
                return
            para = doc.add_paragraph()
            r = para.add_run(f"{label}: ")
            r.bold = True
            para.add_run(value)

        field("Hook", item.get("hook", ""))
        field("Format", item.get("format", ""))
        field("Pillar", item.get("pillar", ""))
        if item.get("outline"):
            para = doc.add_paragraph()
            para.add_run("Outline:").bold = True
            for beat in item["outline"]:
                doc.add_paragraph(beat, style="List Number")
        field("Keywords", ", ".join(item.get("keywords") or []))
        field("Call to action", item.get("cta", ""))
        tags = " · ".join(x for x in (item.get("goal"), item.get("effort")) if x)
        field("Goal / effort", tags)

    if plan.get("assumptions"):
        doc.add_heading("Assumptions", level=1)
        doc.add_paragraph(
            "The agent assumed the following because it wasn't told otherwise:"
        )
        for a in plan["assumptions"]:
            doc.add_paragraph(a, style="List Bullet")

    out = io.BytesIO()
    doc.save(out)
    logger.info("[plan-export] docx: %d item(s)", len(plan.get("items") or []))
    return out.getvalue()


EXPORTERS = {
    "xlsx": (to_xlsx, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "csv": (to_csv, "text/csv; charset=utf-8"),
}


# ---------------------------------------------------------------------------
# Taking ONE SCRIPT out of the app
# ---------------------------------------------------------------------------
# A script leaves two ways, and they are deliberately different documents:
#
#   .txt   — the MACHINE-READABLE one. Byte for byte what
#            `plan_agent.script_to_text` produced, which is byte for byte what
#            the storyboard breakdown reads. A creator who exports this, edits
#            it in a text editor and pastes it back into Script to Storyboard
#            gets exactly the document the app would have handed over itself.
#            ⚠ DO NOT PRETTY-PRINT IT. Every heading, blank line and `NAME:`
#            prefix in there is something the breakdown parses.
#
#   .docx  — the HUMAN one, laid out the way a screenplay is: slug lines, action
#            in the margin, dialogue indented under the speaker's name. This one
#            is for reading, for the client, for the person holding the camera.


def to_script_txt(script: dict) -> bytes:
    """The script exactly as the breakdown reads it. See the note above.

    Falls back to rebuilding the text when a stored script predates the `text`
    field, so an old session still exports rather than downloading an empty
    file — through the SAME function, so the format cannot fork.
    """
    text = str(script.get("text") or "").strip()
    if not text:
        from plan_agent import script_to_text

        text = script_to_text(script).strip()
    logger.info("[plan-export] script txt: %d chars", len(text))
    # UTF-8 with a BOM, for the same reason to_csv uses one: this opens in
    # Notepad and Word as often as in a code editor, and without it a Hindi or
    # Hinglish script full of Devanagari renders as mojibake.
    return ("﻿" + text + "\n").encode("utf-8")


def to_script_docx(script: dict) -> bytes:
    """The script laid out as a screenplay, to read or send on."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    doc.add_heading(str(script.get("title") or "Untitled script"), level=0)

    # The header line states the things a reader checks first, and states the
    # runtime as an ESTIMATE against what was asked for — a script that came
    # back at 90 seconds when 45 were asked for should say so on page one.
    bits = []
    if script.get("seconds"):
        estimated = script.get("estimated_seconds")
        bits.append(
            f"target {script['seconds']}s"
            + (f" · reads at about {estimated}s" if estimated else "")
        )
    if script.get("spoken_words"):
        bits.append(f"{script['spoken_words']} spoken words")
    if script.get("rating"):
        bits.append(f"rated {script['rating']}")
    if script.get("language"):
        bits.append(str(script["language"]))
    if bits:
        run = doc.add_paragraph().add_run(" · ".join(bits))
        run.italic = True
        run.font.size = Pt(9)

    if script.get("logline"):
        para = doc.add_paragraph()
        para.add_run("Logline: ").bold = True
        para.add_run(str(script["logline"]))

    if script.get("characters"):
        doc.add_heading("Cast", level=1)
        for c in script["characters"]:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(str(c.get("name", ""))).bold = True
            if c.get("description"):
                para.add_run(f" — {c['description']}")

    for scene in script.get("scenes") or []:
        heading = doc.add_heading(level=1)
        heading.add_run(
            f"{scene.get('number', 1)}. {str(scene.get('heading', '')).upper()}"
        )
        for beat in scene.get("beats") or []:
            kind = beat.get("type")
            text = str(beat.get("text", ""))
            name = str(beat.get("character", "")).upper()

            if kind in ("dialogue", "vo"):
                # Indented under the speaker, the way a screenplay sets it —
                # the one formatting cue that makes a script skimmable.
                who = doc.add_paragraph()
                who.paragraph_format.left_indent = Inches(1.5)
                who.paragraph_format.space_after = Pt(0)
                who.add_run(name + (" (V.O.)" if kind == "vo" else "")).bold = True
                line = doc.add_paragraph(text)
                line.paragraph_format.left_indent = Inches(1.0)
                line.paragraph_format.right_indent = Inches(1.0)
            elif kind == "text":
                para = doc.add_paragraph()
                para.add_run("ON SCREEN: ").bold = True
                para.add_run(text)
            else:
                doc.add_paragraph(text)

    if script.get("cta"):
        doc.add_heading("Call to action", level=1)
        doc.add_paragraph(str(script["cta"]))

    if script.get("notes"):
        doc.add_heading("Production notes", level=1)
        for n in script["notes"]:
            doc.add_paragraph(str(n), style="List Bullet")

    out = io.BytesIO()
    doc.save(out)
    logger.info("[plan-export] script docx: %d scene(s)", len(script.get("scenes") or []))
    return out.getvalue()


SCRIPT_EXPORTERS = {
    "txt": (to_script_txt, "text/plain; charset=utf-8"),
    "docx": (to_script_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
}
